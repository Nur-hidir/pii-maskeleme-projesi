import re, io, tempfile, os
from dataclasses import dataclass, field
from typing import List
from collections import defaultdict
from functools import lru_cache

import streamlit as st
import wikipediaapi
from transformers import AutoTokenizer, AutoModelForTokenClassification, pipeline
import pdfplumber
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfgen import canvas
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from PIL import Image
import pytesseract
import numpy as np
import cv2
import streamlit.components.v1 as components

# ============================================================
# REGEX DESENLERİ
# ============================================================
REGEX_PATTERNS = {
    "DATE": r"\b\d{4}[-/]\d{1,2}[-/]\d{1,2}\b|\b\d{1,2}[-/]\d{1,2}([-./]\d{2,4})\b",
    "PHONE": r"(?<!\d)(?:(?:\+|00)?90[\s\-]?|0[\s\-]?)?\(?5\d{2}\)?[\s\-]?\d{3}[\s\-]?\d{2}[\s\-]?\d{2}(?!\d)",
    "IBAN": r"\bTR\d{2}[\s]?\d{4}[\s]?\d{4}[\s]?\d{4}[\s]?\d{4}[\s]?\d{4}[\s]?\d{2}\b",
    "CREDIT_CARD": r"\b(?:\d[ -]*?){13,16}\b",
    "EMAIL": r"[a-zA-Z0-9._%+\-çğıöşüÇĞIİÖŞÜ]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}",
    "IP": r"\b\d{1,3}(?:\.\d{1,3}){3}\b",
    "TCKN": r"\b[1-9][0-9]{10}\b",
    "PASSPORT": r"\b(?=.*[A-Z])[A-Z0-9]{6,9}\b",
}
REGEX_ORDER = ["PHONE", "IBAN", "CREDIT_CARD", "EMAIL", "TCKN", "PASSPORT", "IP", "DATE"]

PRIVATE_EMAIL_DOMAINS = ["@gmail.com", "@yahoo.com", "@outlook.com"]

def is_private_email(email: str) -> bool:
    return any(email.lower().endswith(d) for d in PRIVATE_EMAIL_DOMAINS)

# ============================================================
# FONT KAYIT
# ============================================================
def register_fonts():
    font_paths = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"
    ]
    for path in font_paths:
        if os.path.exists(path):
            pdfmetrics.registerFont(TTFont("CustomUnicodeFont", path))
            return "CustomUnicodeFont"
    return "Helvetica"

DEFAULT_FONT = register_fonts()

# ============================================================
# DOSYA OKUMA
# ============================================================
def read_pdf(file_bytes: bytes) -> str:
    text = ""
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def read_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    return "\n".join([p.text for p in doc.paragraphs])

def read_image(file_bytes: bytes) -> str:
    image = Image.open(io.BytesIO(file_bytes)).convert("RGB")
    w, h = image.size
    image = image.resize((w * 2, h * 2), Image.LANCZOS)
    text = pytesseract.image_to_string(
        image,
        lang="tur+eng",
        config="--oem 3 --psm 3"
    )
    return text

# ============================================================
# DATA MODEL
# ============================================================
@dataclass
class Annotation:
    type: str
    value: str
    start: int
    end: int
    source: str
    level: int = None
    action: str = None
    reason: List[str] = field(default_factory=list)
    placeholder: str = None

# ============================================================
# BERT NER
# ============================================================
@st.cache_resource(show_spinner="🤖 NER modeli yükleniyor...")
def load_ner_pipelines():
    models = {
        "tr": "savasy/bert-base-turkish-ner-cased",
    }
    pipes = {}
    for lang, model_name in models.items():
        pipes[lang] = pipeline(
            "ner",
            model=AutoModelForTokenClassification.from_pretrained(model_name),
            tokenizer=AutoTokenizer.from_pretrained(model_name),
            aggregation_strategy="max",
        )
    return pipes

def predict_bert(text: str, lang: str, pipes: dict) -> List[Annotation]:
    return [
        Annotation(
            type=e["entity_group"],
            value=e["word"],
            start=e["start"],
            end=e["end"],
            source="bert",
        )
        for e in pipes[lang](text)
    ]

# ============================================================
# REGEX TESPİTİ
# ============================================================
def find_regex_entities(text: str) -> List[Annotation]:
    out = []
    for label in REGEX_ORDER:
        for m in re.finditer(REGEX_PATTERNS[label], text, re.MULTILINE | re.DOTALL):
            out.append(Annotation(label, m.group(), m.start(), m.end(), "regex"))
    return out

# ============================================================
# ENTITY MERGE
# ============================================================
def merge_entities(a: List[Annotation], b: List[Annotation]) -> List[Annotation]:
    priority = {"TCKN": 5, "PHONE": 3, "IBAN": 6, "CREDIT_CARD": 5, "EMAIL": 4, "DATE": 2, "IP": 3}
    all_e = sorted(a + b, key=lambda x: (x.start, -(x.end - x.start)))
    merged = []
    for e in all_e:
        if not merged:
            merged.append(e)
        elif e.start >= merged[-1].end:
            merged.append(e)
        elif priority.get(e.type.upper(), 2) > priority.get(merged[-1].type.upper(), 1):
            merged[-1] = e
    return merged

# ============================================================
# WIKIPEDIA PUBLIC ENTITY CHECK
# ============================================================
@lru_cache(maxsize=1000)
def is_public_entity(name: str, lang: str) -> bool:
    clean_name = re.sub(r"[''].*$", "", name).strip()
    try:
        wiki = wikipediaapi.Wikipedia(
            user_agent="PII_Masking_App/1.0 (contact@example.com)",
            language=lang,
            extract_format=wikipediaapi.ExtractFormat.WIKI,
        )
        return wiki.page(clean_name).exists()
    except Exception:
        return False

# ============================================================
# CONTEXTUAL SCORING (ONLY KVKK)
# ============================================================
def contextual_scoring(entity: Annotation, text: str, lang: str):
    et = entity.type.upper()
    if et == "PER":
        et = "PERSON"
    if et == "LOC":
        et = "LOCATION"
    entity.type = et
    val = entity.value

    if et in ("PERSON", "LOCATION", "ORG") and val and val[0].islower():
        return 3, "PASS", ["bert_false_positive_lowercase"]
    if et in ("PERSON", "ORG", "LOCATION") and is_public_entity(val, lang):
        return 3, "PASS", ["public_entity_wikipedia"]

    # Sadece KVKK mantığı
    if et in ("TCKN", "CREDIT_CARD", "PASSPORT", "IBAN"):
        return 1, "MASK", ["kvkk_special_category"]
    if et == "EMAIL":
        if is_private_email(val):
            return 2, "MASK", ["private_email_domain"]
        return 3, "PASS", ["non_private_email"]
    if et in ("PERSON", "PHONE", "DATE", "IP", "LOCATION", "ORG"):
        return 2, "MASK", ["kvkk_personal_data"]
    
    return 3, "PASS", ["kvkk_anonymous"]

# ============================================================
# MASKING
# ============================================================
def apply_placeholder_masking(text: str, annotations: List[Annotation]):
    counters = defaultdict(int)
    mapping = {}
    display_labels = {
        "PHONE": "TEL_NUMBER", "DATE": "DATE",
        "PERSON": "PERSON", "LOCATION": "LOCATION", "ORG": "ORGANIZATION",
    }
    for a in annotations:
        if a.action != "MASK":
            continue
        key = (a.type.upper(), a.value)
        if key not in mapping:
            base_label = display_labels.get(a.type.upper(), a.type.upper())
            counters[base_label] += 1
            mapping[key] = f"[{base_label}_{counters[base_label]}]"
        a.placeholder = mapping[key]

    for a in sorted([x for x in annotations if x.action == "MASK"], key=lambda x: x.start, reverse=True):
        text = text[:a.start] + a.placeholder + text[a.end:]
    return text, mapping

# ============================================================
# EXPORT
# ============================================================
def export_masked_pdf(text: str) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4)
    styles = getSampleStyleSheet()
    style = styles["Normal"]
    style.fontName = DEFAULT_FONT
    style.fontSize = 11
    style.leading = 14
    story = []
    for line in text.split("\n"):
        if line.strip() == "":
            story.append(Spacer(1, 8))
        else:
            story.append(Paragraph(line, style))
    doc.build(story)
    return buf.getvalue()

def export_masked_docx(text: str) -> bytes:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    for line in text.split("\n"):
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ============================================================
# ANA ANALİZ FONKSİYONU
# ============================================================
def analyze_text(text: str, lang: str, pipes: dict):
    regex_e = find_regex_entities(text)
    bert_e = predict_bert(text, lang, pipes)
    entities = merge_entities(regex_e, bert_e)
    for e in entities:
        lvl, act, rsn = contextual_scoring(e, text, lang)
        e.level = lvl
        e.action = act
        e.reason = rsn
    masked_text, mapping = apply_placeholder_masking(text, entities)
    return entities, masked_text, mapping

# ============================================================
# STREAMLIT UI
# ============================================================
st.set_page_config(
    page_title="PII Masking Studio",
    page_icon="🛡️",
    layout="wide",
)

# Dil seçimi kaldırıldı, varsayılan Türkçe (tr) olarak sabitlendi
lang = "tr"




# ---------- CUSTOM CSS ----------
st.markdown("""
<style>
/* BACKGROUND */
body {
    background: radial-gradient(circle at 20% 20%, #0f172a, #020617);
}
body::before {
    content: "";
    position: fixed;
    inset: 0;
    background-image: 
        linear-gradient(rgba(255,255,255,0.03) 1px, transparent 1px),
        linear-gradient(90deg, rgba(255,255,255,0.03) 1px, transparent 1px);
    background-size: 40px 40px;
    pointer-events: none;
}
body::after {
    content: "";
    position: fixed;
    inset: 0;
    background: radial-gradient(circle at 80% 20%, rgba(0,255,200,0.08), transparent 40%);
    pointer-events: none;
}
/* HEADER */
.title {
    font-size: 34px;
    font-weight: 700;
    color: white;
}
.subtitle {
    color: #94a3b8;
    margin-bottom: 25px;
}
/* CARD */
.input-card {
    text-align: center;
    padding: 25px;
    border-radius: 18px;
    background: rgba(255,255,255,0.04);
    transition: 0.25s;
}
.input-card:hover {
    transform: translateY(-6px) scale(1.03);
    background: rgba(255,255,255,0.08);
    box-shadow: 0 0 25px rgba(0,255,200,0.2);
}
/* ICON */
.icon {
    width: 36px;
    height: 36px;
    margin-bottom: 10px;
    color: #94a3b8;
}
.input-card:hover .icon {
    color: #00ffc8;
}
/* BUTTON */
.stButton>button {
    border-radius: 12px;
    height: 50px;
    font-weight: 600;
    background: linear-gradient(90deg, #00ffc8, #00c3ff);
    color: black;
    border: none;
}
/* TEXT AREA */
textarea {
    background: rgba(255,255,255,0.05) !important;
    color: white !important;
    border-radius: 12px !important;
}
/* METRICS */
[data-testid="stMetric"] {
    background: rgba(255,255,255,0.04);
    padding: 15px;
    border-radius: 12px;
}

</style>
""", unsafe_allow_html=True)



# ---------- HEADER ----------
st.markdown("""
<div class="title">PII Masking Studio</div>
<div class="subtitle">Real-time AI data anonymization engine</div>
""", unsafe_allow_html=True)



# ---------- MODEL LOAD ----------
pipes = load_ner_pipelines()

# ---------- INPUT MODE ----------
st.markdown("### Select Input Type")

if "input_mode" not in st.session_state:
    st.session_state.input_mode = "text"

col1, col2, col3 = st.columns(3)

# --- TEXT ---
with col1:
    st.markdown("""
    <div class="input-card">
        <svg class="icon" viewBox="0 0 24 24" fill="none" stroke="currentColor">
            <path d="M4 7h16M4 12h16M4 17h10"/>
        </svg>
        <b>Text</b>
    </div>
    """, unsafe_allow_html=True)

    if st.button("Select Text", use_container_width=True):
        st.session_state.input_mode = "text"

# --- DOCUMENT ---
with col2:
    st.markdown("""
    <div class="input-card">
        <svg class="icon" viewBox="0 0 24 24" fill="none" stroke="currentColor">
            <path d="M6 2h9l5 5v15H6z"/>
        </svg>
        <b>Document</b>
    </div>
    """, unsafe_allow_html=True)

    if st.button("Select Document", use_container_width=True):
        st.session_state.input_mode = "file"

# --- IMAGE ---
with col3:
    st.markdown("""
    <div class="input-card">
        <svg class="icon" viewBox="0 0 24 24" fill="none" stroke="currentColor">
            <rect x="3" y="3" width="18" height="18"/>
            <circle cx="8" cy="8" r="2"/>
        </svg>
        <b>Image</b>
    </div>
    """, unsafe_allow_html=True)

    if st.button("Select Image", use_container_width=True):
        st.session_state.input_mode = "image"

mode = st.session_state.input_mode

st.markdown("---")

input_text = ""
file_type = None

# ---------- INPUT AREA ----------
if mode == "text":
    input_text = st.text_area("", height=200, placeholder="Enter your text...")

elif mode == "file":
    uploaded_file = st.file_uploader("", type=["pdf", "docx"])
    if uploaded_file:
        file_bytes = uploaded_file.read()
        if uploaded_file.name.endswith(".pdf"):
            input_text = read_pdf(file_bytes)
        else:
            input_text = read_docx(file_bytes)

elif mode == "image":
    uploaded_image = st.file_uploader("", type=["png", "jpg", "jpeg"])
    if uploaded_image:
        st.image(uploaded_image, use_column_width=True)
        with st.spinner("Running OCR..."):
            input_text = read_image(uploaded_image.read())

st.markdown("---")

# ---------- RUN ----------
run = st.button("🚀 Analyze & Mask Data", use_container_width=True)

# ---------- RESULTS ----------
if run:
    if not input_text.strip():
        st.warning("Please provide input.")
    else:
        with st.spinner("Analyzing..."):
            entities, masked_text, mapping = analyze_text(input_text, lang, pipes)

        st.markdown("### 📊 Results")

        masked_count = sum(1 for e in entities if e.action == "MASK")
        passed_count = sum(1 for e in entities if e.action == "PASS")

        col1, col2, col3 = st.columns(3)
        col1.metric("Entities", len(entities))
        col2.metric("Masked", masked_count)
        col3.metric("Passed", passed_count)

        st.markdown("---")

        left, right = st.columns([1.3, 1])

        with left:
            st.text_area("Masked Output", value=masked_text, height=300)

        with right:
            import pandas as pd
            df = pd.DataFrame([{
                "Type": e.type,
                "Value": e.value,
                "Action": e.action,
                "Level": e.level
            } for e in entities])
            st.dataframe(df, use_container_width=True)

        
        st.markdown("<div style='margin-top:-20px'></div>", unsafe_allow_html=True)
        c1, c2 = st.columns([1, 1])

        with c1:
            st.download_button(
                "Download PDF",
                export_masked_pdf(masked_text),
                file_name="masked.pdf",
                use_container_width=True
            )

        with c2:
            st.download_button(
                "Download DOCX",
                export_masked_docx(masked_text),
                file_name="masked.docx",
                use_container_width=True
            )
